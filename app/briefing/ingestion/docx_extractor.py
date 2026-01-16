"""
DOCX Text Extractor

Extracts text from DOCX/Word files stored in Replit Object Storage.
Handles large files efficiently using temporary files when needed.
"""

import logging
import tempfile
import os
from typing import Optional
from io import BytesIO

logger = logging.getLogger(__name__)

# File size limits
MAX_DOCX_SIZE_BYTES = 50 * 1024 * 1024  # 50 MB max
LARGE_FILE_THRESHOLD = 10 * 1024 * 1024  # 10 MB - use temp file above this
MAX_TEXT_LENGTH = 5 * 1024 * 1024  # 5 MB max extracted text


def extract_text_from_docx(storage_key: str) -> Optional[str]:
    """
    Extract text from a DOCX file in Replit Object Storage.

    Handles large files efficiently:
    - Files under 10MB: Process in memory
    - Files 10-50MB: Use temporary file
    - Files over 50MB: Rejected

    Args:
        storage_key: Replit Object Storage key (e.g., 'uploads/user_123/document.docx')

    Returns:
        Extracted text as string, or None if extraction fails
    """
    try:
        from replit.object_storage import Client
        client = Client()

        # Download DOCX from storage
        docx_data = client.download_bytes(storage_key)
        file_size = len(docx_data)

        logger.info(f"Processing DOCX {storage_key} ({file_size / 1024 / 1024:.2f} MB)")

        # Check size limit
        if file_size > MAX_DOCX_SIZE_BYTES:
            logger.warning(f"DOCX too large: {file_size} bytes exceeds {MAX_DOCX_SIZE_BYTES} byte limit")
            return None

        # Decide processing strategy based on size
        if file_size > LARGE_FILE_THRESHOLD:
            logger.info(f"Using temp file for large DOCX ({file_size / 1024 / 1024:.2f} MB)")
            return _extract_with_tempfile(docx_data, storage_key)
        else:
            return _extract_in_memory(docx_data, storage_key)

    except Exception as e:
        logger.error(f"Error extracting text from DOCX {storage_key}: {e}", exc_info=True)
        return None


def _extract_in_memory(docx_data: bytes, storage_key: str) -> Optional[str]:
    """Extract text from DOCX data in memory."""
    try:
        from docx import Document

        docx_file = BytesIO(docx_data)
        doc = Document(docx_file)

        text_parts = []
        total_length = 0

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text = paragraph.text.strip()
                text_parts.append(text)
                total_length += len(text)

                if total_length > MAX_TEXT_LENGTH:
                    logger.warning(f"Extracted text exceeds {MAX_TEXT_LENGTH} chars, truncating")
                    break

        # Extract text from tables if we haven't hit the limit
        if total_length < MAX_TEXT_LENGTH:
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text = ' | '.join(row_text)
                        text_parts.append(text)
                        total_length += len(text)

                        if total_length > MAX_TEXT_LENGTH:
                            logger.warning(f"Extracted text exceeds limit in tables, truncating")
                            break
                if total_length > MAX_TEXT_LENGTH:
                    break

        extracted_text = '\n\n'.join(text_parts)

        if not extracted_text or len(extracted_text.strip()) < 50:
            logger.warning(f"DOCX extraction returned minimal text for {storage_key}")
            return None

        logger.info(f"Extracted {len(extracted_text)} chars from DOCX {storage_key}")
        return extracted_text.strip()[:MAX_TEXT_LENGTH]

    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return None


def _extract_with_tempfile(docx_data: bytes, storage_key: str) -> Optional[str]:
    """Extract text using a temporary file for large DOCX files."""
    temp_path = None

    try:
        from docx import Document

        # Write to temp file
        fd, temp_path = tempfile.mkstemp(suffix='.docx')
        try:
            os.write(fd, docx_data)
        finally:
            os.close(fd)

        # Clear in-memory data to free memory
        del docx_data

        # Open from file
        doc = Document(temp_path)

        text_parts = []
        total_length = 0

        # Extract text from paragraphs
        para_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text = paragraph.text.strip()
                text_parts.append(text)
                total_length += len(text)
                para_count += 1

                if total_length > MAX_TEXT_LENGTH:
                    logger.warning(f"Extracted text exceeds limit at paragraph {para_count}")
                    break

                # Log progress for large files
                if para_count > 0 and para_count % 100 == 0:
                    logger.info(f"Processed {para_count} paragraphs...")

        # Extract text from tables if we haven't hit the limit
        if total_length < MAX_TEXT_LENGTH:
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text = ' | '.join(row_text)
                        text_parts.append(text)
                        total_length += len(text)
                        table_count += 1

                        if total_length > MAX_TEXT_LENGTH:
                            break
                if total_length > MAX_TEXT_LENGTH:
                    break

            if table_count > 0:
                logger.info(f"Extracted {table_count} table rows")

        extracted_text = '\n\n'.join(text_parts)

        if not extracted_text or len(extracted_text.strip()) < 50:
            logger.warning(f"DOCX extraction returned minimal text for {storage_key}")
            return None

        logger.info(f"Extracted {len(extracted_text)} chars from large DOCX {storage_key}")
        return extracted_text.strip()[:MAX_TEXT_LENGTH]

    except ImportError:
        logger.error("python-docx not installed for temp file extraction")
        return None

    finally:
        # Clean up temp file
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception as e:
                logger.warning(f"Could not delete temp file {temp_path}: {e}")
